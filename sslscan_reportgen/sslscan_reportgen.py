#!/usr/bin/env python

"""script to parse output from sslscan and find common issues, then dump into a docx"""

#todo: strip colorized output

try:
    import argparse
    import docx
    import os
    import re
    import sys
    import time
    from docx.shared import Pt
    from docx.shared import RGBColor
except Exception as e:
    print('\n[!] Import(s) failed!: {}'.format(e))

class sslscan_beautifier():
    """a class is probably unnecessary for this script"""
    def __init__(self, args):
        #defaults
        #pass in args. this is messy
        self.args = args
        #start time
        self.start_time = time.time()
        #verbosity explicitly off
        self.verbose = False
        #version
        self.version = 'beta.09_15_2017'
        #file to read
        self.ssl_file = None
        #misconfiguration dictionaries
        self.rc4_dict = {}
        self.sslv2_dict = {}
        self.sslv3_dict = {}
        self.des_dict = {}
        self.tls10_dict = {}
        self.weakbits_dict = {}
        self.heartbleed_dict = {}
        self.md5_dict = {}
        #store sslscan host results here for parsing key, val is ip, results
        self.result_dictionary = {}
        #dump reports here
        self.report_dir = './reports/'
        #check for report directory, make it if not found
        if not os.path.exists(self.report_dir):
            os.makedirs(self.report_dir)
        #show user script has started
        print('\n**** SSLScan ReportGen ****\n')

    def check_args(self, parser):
        """Validates supplied args"""
        if self.args.verbose is True: print(\
            '[i] Version: {}\n[i] Options: {}'.format(self.version, parser.parse_args()))
        #make sure -f is supplied
        if not self.args.file:
            print ('No file supplied')
            parser.print_help()
            sys.exit(0)
        #make sure -c is set for file naming purposes
        if not self.args.client:
            print ('Please provide a client name')
            parser.print_help()
            sys.exit(0)
        #clean up client name, remove spaces and stuff
        self.args.client = ''.join(e for e in self.args.client if e.isalnum())

    def read_file(self):
        """open targets file, splits on ip and puts ip and the result into a dictionary"""
        try:
            with open(self.args.file) as f:
                #use a regex to split the file into sections, delimited by the word Testing
                for i, result in enumerate(re.findall('Testing(.*?)Testing', f.read(), re.S)):
                    #look for the first line from the Testing
                    # BUG for some reason the regex removes 'testing' from the result
                    if 'SSL server' in result:
                        #set up a dictionary key of ip, val of the result
                        self.result_dictionary\
                        [re.findall(r'(?:\d{1,3}\.)+(?:\d{1,3})', result)[0]] = result
        except Exception as e:
            print('\n[!] Couldn\'t open file: {}'.format(e))

    def parse_dict(self):
        """reads thru result_dictionary and search for misconfiguraitons in the scan results"""
        #run through the dictonary containing IPs and the scan results as keys
        for key, val in self.result_dictionary.items():
            #for each line of the scan output
            #print(key)
            for line in val.splitlines():
                #look for RC4
                if 'RC4' in line:
                    self.rc4_dict[key] = line
                if 'SSLv3' in line:
                    self.sslv3_dict[key] = line
                if 'SSLv2' in line:
                    self.sslv2_dict[key] = line
                if 'MD5' in line:
                    self.md5_dict[key] = line
                if 'DES' in line:
                    self.des_dict[key] = line
                if 'TLSv1.0' in line:
                    self.tls10_dict[key] = line
                if 'bits' in line:
                    bits = line.split()[2]
                    if int(bits) < 128:
                        self.weakbits_dict[key] = line
                if 'to heartbleed' in line:
                    if not 'not' in line:
                        self.heartbleed_dict[key] = line

    def gen_report(self):
        """create a document -- to use a template put the docx in the (), like (my_template.docx)"""
        document = docx.Document()
        print ('\n**** Generating Report: ****')
        print('**** {}sslscan_{}.docx ****\n'.format(self.report_dir, self.args.client))
        #optional header
        heading = document.add_heading(level=3)
        run_heading = heading.add_run('Hosts With Weak Transport Security')
        font = run_heading.font
        font.name = 'Arial'
        font.size = Pt(24)
        font.color.rgb = RGBColor(0xe9, 0x58, 0x23)
        paragraph = document.add_paragraph()
        #start by printing the ip address in 20pt arial
        #for each ip parsed out, look thru the results and add them to the report
        for key, val in sorted(self.result_dictionary.items()):
            run_paragraph = paragraph.add_run(key)
            font = run_paragraph.font
            font.name = 'Arial'
            font.size = Pt(20)
            run_paragraph = paragraph.add_run('\n')
            if self.args.verbose is True:
                print('Added host: {}'.format(key))
            #add RC4 findings
            if self.rc4_dict:
                run_paragraph = paragraph.add_run('RC4 Support\n')
                font = run_paragraph.font
                font.name = 'Arial'
                font.size = Pt(16)
                for i, r in self.rc4_dict.items():
                    run_paragraph = paragraph.add_run(r)
                    font = run_paragraph.font
                    font.name = 'Arial'
                    font.size = Pt(11)
                    run_paragraph = paragraph.add_run('\n')
            if self.sslv2_dict:
                #if this host had sslv2, print the header
                run_paragraph = paragraph.add_run('SSLv2 Support\n')
                font = run_paragraph.font
                font.name = 'Arial'
                font.size = Pt(16)
                for i, r in self.sslv2_dict.items():
                    run_paragraph = paragraph.add_run(r)
                    font = run_paragraph.font
                    font.name = 'Arial'
                    font.size = Pt(11)
                    run_paragraph = paragraph.add_run('\n')
            if self.sslv3_dict:
                #add sslv3 finding
                run_paragraph = paragraph.add_run('SSLv3 Support\n')
                font = run_paragraph.font
                font.name = 'Arial'
                font.size = Pt(16)
                for i, r in self.sslv3_dict.items():
                    run_paragraph = paragraph.add_run(r)
                    font = run_paragraph.font
                    font.name = 'Arial'
                    font.size = Pt(11)
                    run_paragraph = paragraph.add_run('\n')
            if self.des_dict:
                #add des finding
                run_paragraph = paragraph.add_run('DES Support\n')
                font = run_paragraph.font
                font.name = 'Arial'
                font.size = Pt(16)
                for i, r in self.des_dict.items():
                    run_paragraph = paragraph.add_run(r)
                    font = run_paragraph.font
                    font.name = 'Arial'
                    font.size = Pt(11)
                    run_paragraph = paragraph.add_run('\n')
            if self.tls10_dict:
                #add tls 1.0 findings
                run_paragraph = paragraph.add_run('TLS v1.0\n')
                font = run_paragraph.font
                font.name = 'Arial'
                font.size = Pt(16)
                for i, r in self.tls10_dict.items():
                    run_paragraph = paragraph.add_run(r)
                    font = run_paragraph.font
                    font.name = 'Arial'
                    font.size = Pt(11)
                    run_paragraph = paragraph.add_run('\n')
            if self.weakbits_dict:
                #add weak key findings
                run_paragraph = paragraph.add_run('Weak key size \n')
                font = run_paragraph.font
                font.name = 'Arial'
                font.size = Pt(16)
                for i, r in self.weakbits_dict.items():
                    run_paragraph = paragraph.add_run(r)
                    font = run_paragraph.font
                    font.name = 'Arial'
                    font.size = Pt(11)
                    run_paragraph = paragraph.add_run('\n')
            if self.heartbleed_dict:
                #add heartbleed findings
                run_paragraph = paragraph.add_run('Heartbleed \n')
                font = run_paragraph.font
                font.name = 'Arial'
                font.size = Pt(16)
                for i, r in self.heartbleed_dict.items():
                    run_paragraph = paragraph.add_run(r)
                    font = run_paragraph.font
                    font.name = 'Arial'
                    font.size = Pt(11)
                    run_paragraph = paragraph.add_run('\n')
            if self.md5_dict:
                #add MD5 findings
                run_paragraph = paragraph.add_run('MD5 Supported \n')
                font = run_paragraph.font
                font.name = 'Arial'
                font.size = Pt(16)
                #enumerate the lines were rc4 was found
                for i, r in self.md5_dict.items():
                    run_paragraph = paragraph.add_run(r)
                    font = run_paragraph.font
                    font.name = 'Arial'
                    font.size = Pt(11)
                    run_paragraph = paragraph.add_run('\n')
        document.save('{}sslscan_{}.docx'.format(self.report_dir, self.args.client))

    def print_summary(self):
        """displays what was found to terminal"""
        if self.rc4_dict:
            print('\n********RC4 Hosts********')
            for k, v in sorted(self.rc4_dict.items()):
                print('{}'.format(k))
        if self.sslv2_dict:
            print('\n********SSLv2 Hosts********')
            for k, v in sorted(self.sslv2_dict.items()):
                print('{}'.format(k))
        if self.sslv3_dict:
            print('\n********SSLv3 Hosts********')
            for k, v in sorted(self.sslv3_dict.items()):
                print('{}'.format(k))
        if self.des_dict:
            print('\n********DES Hosts********')
            for k, v in sorted(self.des_dict.items()):
                print('{}'.format(k))
        if self.tls10_dict:
            print('\n********TLS v1.0 Hosts********')
            for k, v in sorted(self.tls10_dict.items()):
                print('{}'.format(k))
        if self.weakbits_dict:
            print('\n********Weak Key Size********')
            for k, v in sorted(self.weakbits_dict.items()):
                print('{}'.format(k))
        if self.heartbleed_dict:
            print('\n********Heartbleed Hosts********')
            for k, v in sorted(self.heartbleed_dict.items()):
                print('{}'.format(k))
        if self.md5_dict:
            print('\n********MD5 Hosts********')
            for k, v in sorted(self.md5_dict.items()):
                print('{}'.format(k))

    def end(self):
        """ending stuff, right now just shows how long script took to run"""
        print('\nCompleted in {:.2f} seconds\n'.format(time.time() - self.start_time))

def main():
    #gather options
    parser = argparse.ArgumentParser()
    parser.add_argument('-c', '--client', metavar='<client>', help='client name')
    parser.add_argument('-f', '--file', metavar='<file>', help='file to read')
    parser.add_argument('-v', '--verbose', help='Optionally enable verbosity', action='store_true')
    args = parser.parse_args()
    run = sslscan_beautifier(args)
    run.check_args(parser)
    run.read_file()
    run.parse_dict()
    run.print_summary()
    run.gen_report()
    run.end()

if __name__ == '__main__':
    main()
